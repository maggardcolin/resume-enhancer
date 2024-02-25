# author: Nick Pastore
import math

from ResumeObjects import Activity, Education, Experience, Project
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# Default 'style' for resume
name_size = 24
body_size = 13
section_header_size = 17
section_space_size = 5

name_pt = Pt(name_size)
body_pt = Pt(body_size)
section_header_pt = Pt(section_header_size)
section_space_pt = Pt(section_space_size)

top_bottom_margin = 1
left_right_margin = 1

font = 'Calibri'


# horizontal line function borrowed from stackoverflow
# https://stackoverflow.com/questions/39006878/python-docx-add-horizontal-line
def insert_hr(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
                              'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                              'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                              'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                              'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                              'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                              'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                              'w:pPrChange'
                              )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)


def format_run(run, pt, bold=False, italic=False):
    """
    Quick helper method to format paragraph runs easier
    :param run: The run you would like to format
    :param pt: Font size you want the run to be
    :param bold: if True make run bold. Default: False
    :param italic: if True make run italic. Default: False
    """
    run.font.name = font
    run.font.size = pt
    run.font.bold = bold
    run.font.italic = italic


class Resume:
    def __init__(self, name, city, state, email, number, objective, links):
        # string
        self.name = name
        # string
        self.email = email
        # string
        self.number = number
        # list of strings
        self.links = links
        # strings
        self.city = city
        # strings
        self.state = state
        # strings
        self.objective = objective
        self.activities = []
        self.experiences = []
        self.educations = []
        self.projects = []
        self.skills = ""

    def add_education(self, degree, date, college, location, gpa, description, delimiter="\n"):
        """
        Adds an education field to the resume
        :param degree: Degree and Major to be added in Education (Ex: Bachelor of Science in Computer Science)
        :param date: Month and year of (expected) graduation (Ex: August 2023)
        :param location: The location where the Education was recieved (Ex: University of Wisconsin - Madison)
        :param gpa: The GPA at the establishment (Ex: 3.87/4.00)
        :param description: Additional fields that will be displayed as bullets. Seperated by a configurable delimiter.
        (commas by default)
        :param delimiter: The delimiter you would like to use for seperate bullet points. Default: '\n'
        """
        e1 = Education.Education(degree, date, college, location, gpa, description, delimiter)
        self.educations.append(e1)

    def add_activity(self, organization, location, role_title, description, delimiter="\n"):
        """
        Adds an activity field to the resume
        :param organization: The organization affiliated with the activity
        :param location: the location the activity took place in
        :param role_title: a short title that best describes your role in the activity
        :param description: Additional fields that will be displayed as bullets. Seperated by a configurable delimiter.
        (commas by default)
        :param delimiter: The delimiter you would like to use for seperate bullet points. Default: '\n'
        """
        a1 = Activity.Activity(organization, location, role_title, description, delimiter)
        self.activities.append(a1)

    def add_experience(self, company, role_title, location, duration, description, delimiter="\n"):
        """
        Adds an (work) experience field to the resume
        :param company: The company you worked for
        :param role_title: The position you've held
        :param location: The location of where you worked
        :param duration: The time period you worked for (Ex: May 2023 - Present)
        :param description: Additional fields that will be displayed as bullets. Seperated by a configurable delimiter.
        (commas by default)
        :param delimiter: The delimiter you would like to use for seperate bullet points. Default: '\n'
        """
        e1 = Experience.Experience(company, role_title, location, duration, description, delimiter)
        self.experiences.append(e1)

    def add_project(self, name, languages_used, description, delimiter="\n"):
        """
        Adds a project field to the resume
        :param name: The name of your project
        :param languages_used: Languages/Tools/Concepts you've used in your project (Ex: 'Java, Python, SQL')
        :param description: Additional fields that will be displayed as bullets. Seperated by a configurable delimiter.
        (commas by default)
        :param delimiter: The delimiter you would like to use for seperate bullet points. Default: '\n'
        """
        p1 = Project.Project(name, languages_used, description, delimiter)
        self.projects.append(p1)

    def add_skills(self, skill):
        """
        Adds skills to your resume
        :param skill: Skills you would like to add. (Ex: 'Java' or 'Java,Python,HTML')
        """
        if len(self.skills) == 0:
            self.skills = skill
        else:
            self.skills = self.skills + "," + skill

    def compile_resume(self):
        """
        compiles the instance variables to generate a docx file using the provided information
        :return: a document object
        """
        # instantiate the doc
        doc = Document()
        verticle_height = self.calculate_vertical_pt_sum()
        global name_size, body_size, section_header_size, left_right_margin, top_bottom_margin, name_pt, body_pt
        global section_header_pt
        while verticle_height >= 1:
            print(verticle_height := self.calculate_vertical_pt_sum())
            name_size -= 0.5
            body_size -= 0.5
            section_header_size -= 0.5
            left_right_margin -= 0.05
            top_bottom_margin -= 0.05
        name_pt = Pt(name_size)
        body_pt = Pt(body_size)
        section_header_pt = Pt(section_header_size)
        # generates the name header
        namep = doc.add_paragraph()
        name = namep.add_run(self.name)
        name.font.name = font
        name.font.size = name_pt
        name.font.bold = True
        namep.paragraph_format.space_after = section_space_pt

        # insert a horizontal line after the Name
        insert_hr(namep)

        # contact "line" after the Name
        contactp = doc.add_paragraph()
        contact_field = self.city + ", " + self.state + " | " + self.email + " | " + self.number
        if len(self.links) != 0:
            contact_field = contact_field + " | " + self.links

        contact = contactp.add_run(contact_field)

        contact.font.name = font
        contact.font.size = body_pt
        contactp.paragraph_format.space_after = section_space_pt

        # if theres an objective statement, Add the header + description
        if len(self.objective) != 0:
            objective_paragraph = doc.add_paragraph()
            objective = objective_paragraph.add_run("Objective")
            format_run(objective, section_header_pt, bold=True)
            objective_paragraph.paragraph_format.space_after = Pt(0)

            objective_description_paragraph = doc.add_paragraph()
            objective_description = objective_description_paragraph.add_run(self.objective)
            objective_description_paragraph.paragraph_format.space_after = section_space_pt
            format_run(objective_description, body_pt)

        # if there are educations added, add the header + fields
        if len(self.educations) != 0:
            education_header_paragraph = doc.add_paragraph()
            education_header = education_header_paragraph.add_run("Education")
            format_run(education_header, section_header_pt, bold=True)
            education_header_paragraph.paragraph_format.space_after = Pt(0)
            for education in self.educations:
                education.import_to_doc(doc, body_pt)

        # if there are experiences added, add the header and the fields
        if len(self.experiences) != 0:
            experience_header_paragraph = doc.add_paragraph()
            experience_header_paragraph.paragraph_format.space_before = section_space_pt
            experience_header_paragraph.paragraph_format.space_after = Pt(0)
            experience_header = experience_header_paragraph.add_run("Experience")
            format_run(experience_header, section_header_pt, bold=True)
            for exp in self.experiences:
                exp.import_to_doc(doc, body_pt)

        # if activities add header and fields
        if len(self.activities) != 0:
            activity_header_paragraph = doc.add_paragraph()
            activity_header_paragraph.paragraph_format.space_after = Pt(0)
            activity_header_paragraph.paragraph_format.space_before = section_space_pt
            header = "Activities"
            if len(self.activities) == 1:
                header = "Activity"
            activity_header = activity_header_paragraph.add_run(header)
            format_run(activity_header, section_header_pt, bold=True)
            for activity in self.activities:
                activity.import_to_doc(doc, body_pt)

        # if projects add header + fields
        if len(self.projects) != 0:
            project_header_paragraph = doc.add_paragraph()
            project_header_paragraph.paragraph_format.space_before = section_space_pt
            project_header_paragraph.paragraph_format.space_after = Pt(0)
            project_header = project_header_paragraph.add_run("Projects")
            format_run(project_header, section_header_pt, bold=True)
            for project in self.projects:
                project.import_to_doc(doc, body_pt)

        # if there are skills (i hope there are) add the header + field
        if len(self.skills) != 0:
            skill_paragraph = doc.add_paragraph()
            skill_paragraph.paragraph_format.space_before = section_space_pt
            skill_paragraph.paragraph_format.space_after = Pt(0)
            skill_header = skill_paragraph.add_run("Skills")
            format_run(skill_header, section_header_pt, bold=True)
            skills_list = ""
            first = True
            for skill in self.skills.split(","):
                if first:
                    skills_list = skill.strip()
                    first = False
                    continue
                skills_list = skills_list + ", " + skill.strip()
            skill_list_run = doc.add_paragraph().add_run(skills_list)
            format_run(skill_list_run, body_pt)

            # adjust the margins to fit the parameters described at the top of file
        section = doc.sections[-1]
        section.left_margin, section.right_margin = (Inches(left_right_margin), Inches(left_right_margin))
        section.top_margin, section.bottom_margin = (Inches(top_bottom_margin), Inches(top_bottom_margin))

        return doc

    def clone(self):
        """
        Creates a clone of the Resume that's being called on
        :return: a duplicate of self
        """
        # create a new resume that has all the same contructor parameters
        duplicate_docx = Resume(self.name, self.city, self.state, self.email, self.number, self.objective, self.links)

        # for every education go grab its description, put it back into a new string so you can pass it into duplicate
        # then pass in all the self variables again to duplicate
        for education in self.educations:
            description = ""
            for i in range(len(education.description)):
                description += education.description[i]
                if i < len(education.description) - 1:
                    description += "\n"
            duplicate_docx.add_education(education.degree, education.date, education.college, education.location,
                                         education.gpa,
                                         description, "\n")

        # same thing but for experiences
        for experience in self.experiences:
            description = ""
            for i in range(len(experience.description)):
                description += experience.description[i]
                if i < len(experience.description) - 1:
                    description += "\n"
            duplicate_docx.add_experience(experience.company, experience.role_title, experience.location,
                                          experience.duration,
                                          description, "\n")
        # same thing but for activities
        for activity in self.activities:
            description = ""
            for i in range(len(activity.description)):
                description += activity.description[i]
                if i < len(activity.description) - 1:
                    description += "\n"
            duplicate_docx.add_activity(activity.organization, activity.location, activity.role_title, description,
                                        "\n")
        # same thing but for projects
        for project in self.projects:
            description = ""
            for i in range(len(project.description)):
                description += project.description[i]
                if i < len(project.description) - 1:
                    description += "\n"
            duplicate_docx.add_project(project.name, project.languages, description, "\n")
        # pass skills into dupe
        duplicate_docx.add_skills(self.skills)
        return duplicate_docx

    def calculate_vertical_pt_sum(self):
        vertical_pt = 10 + name_size + body_size
        if len(self.objective) != 0:
            vertical_pt += section_header_size
            vertical_pt += math.floor(
                1 + len(self.objective) / (
                        95 * (((8.5 - 2 * left_right_margin) / 8.5) * (13.5 / body_size)))) * body_size
        if len(self.educations) != 0:
            vertical_pt += section_header_size + section_space_size
            for education in self.educations:
                vertical_pt += 3 * body_size
                for desc in education.description:
                    vertical_pt += math.floor(
                        1 + len(desc) / (90 * (((8.5 - 2 * left_right_margin) / 8.5) * (
                                13.5 / body_size)))) * body_size + 0.037 * body_size
        if len(self.experiences) != 0:
            vertical_pt += section_header_size + section_space_size
            for exp in self.experiences:
                vertical_pt += 2 * body_size
                for desc in exp.description:
                    vertical_pt += math.floor(
                        1 + len(desc) / (90 * (((8.5 - 2 * left_right_margin) / 8.5) * (
                                13.5 / body_size)))) * body_size + 0.037 * body_size

        if len(self.activities) != 0:
            vertical_pt += section_header_size + section_space_size
            for activity in self.activities:
                vertical_pt += 2 * body_size
                for desc in activity.description:
                    vertical_pt += math.floor(
                        1 + len(desc) / (90 * (((8.5 - 2 * left_right_margin) / 8.5) * (
                                13.5 / body_size)))) * body_size + 0.037 * body_size

        if len(self.projects) != 0:
            vertical_pt += section_header_size + section_space_size
            for project in self.projects:
                vertical_pt += body_size
                for desc in project.description:
                    vertical_pt += math.floor(
                        1 + len(desc) / (90 * (((8.5 - 2 * left_right_margin) / 8.5) * (
                                13.5 / body_size)))) * body_size + 0.037 * body_size
        if len(self.skills) != 0:
            vertical_pt += section_header_size + section_space_size
            vertical_pt += math.floor(
                1 + len(self.skills) / (90 * (((8.5 - 2 * left_right_margin) / 8.5) * (13.5 / body_size)))) * body_size
        return vertical_pt / (564 * (11 - 2 * top_bottom_margin) / 11)


if __name__ == '__main__':
    res = Resume("bleh", "1", "IL", "EMAIL", "NUMBER",
                 "bfdsfjhdsjkafdhghcjashgsdyifuhjsadkhfiudshfiusdgufhsadfhjkdsahfjksdhjkfhdsakjfhjkdsahfjgscnbgyusihojlkfmnbghjklmnjbvcftgyuiojkmnbvcfdtyuiojknbvcfdrtyuhjnbvcfxdsrtyuihjbvgfcdrt7yuhj",
                 "LinkedIn , GitHub")
    res.add_education("BS in BS", "NOW", "College", "location", "4.00/4.00",
                      "1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,1,11", ",")
    res.compile_resume().save("test.docx")
