# author: Nick Pastore

from docx.shared import Pt

import resume


class Project:
    def __init__(self, name, languages_used, description, delimiter):
        self.name = name
        self.languages = languages_used
        self.description = description.split(delimiter)

    def import_to_doc(self, doc, body_pt):
        """
        Imports this Project to a given document
        :param doc: The document to import to
        :param body_pt: The size to make the body fields
        """
        header_paragraph = doc.add_paragraph()
        header_paragraph.paragraph_format.space_after = Pt(0)

        header_name = header_paragraph.add_run(self.name)
        resume.format_run(header_name, body_pt, bold=True)

        resume.format_run(header_paragraph.add_run(" | "), body_pt)

        header_languages = header_paragraph.add_run(self.languages)
        resume.format_run(header_languages, body_pt, italic=True)

        if len(self.description) != 0:
            for desc in self.description:
                desc_paragraph = doc.add_paragraph()
                desc_paragraph.style = 'List Bullet'
                desc_paragraph.paragraph_format.space_after = Pt(0)
                resume.format_run(desc_paragraph.add_run(desc), body_pt)
