from docx.shared import Pt

import resume


class Experience:
    def __init__(self, company, role_title, location, duration, description, delimiter):
        self.company = company
        self.role_title = role_title
        self.location = location
        self.duration = duration
        self.description = description.split(delimiter)

    def import_to_doc(self, doc, body_pt):
        """
        Imports this Experience to a given document
        :param doc: The document to import to
        :param body_pt: The size to make the body fields
        """
        experience_header_paragraph = doc.add_paragraph()
        experience_header_paragraph.paragraph_format.space_after = Pt(0)

        experience_header_company = experience_header_paragraph.add_run(self.company)
        resume.format_run(experience_header_company, body_pt, bold=True)

        experience_header_seperator = experience_header_paragraph.add_run(" | ")
        resume.format_run(experience_header_seperator, body_pt)

        experience_header_location = experience_header_paragraph.add_run(self.location)
        resume.format_run(experience_header_location, body_pt, italic=True)

        role_title_paragraph = doc.add_paragraph()
        role_title_start = role_title_paragraph.add_run(self.role_title + " | ")
        resume.format_run(role_title_start, body_pt)
        role_title_duration = role_title_paragraph.add_run(self.duration)
        resume.format_run(role_title_duration, body_pt, italic=True)
        role_title_paragraph.paragraph_format.space_after = Pt(0)

        if len(self.description) != 0:
            for desc in self.description:
                desc_paragraph = doc.add_paragraph()
                desc_paragraph.style = 'List Bullet'
                desc_paragraph.paragraph_format.space_after = Pt(0)
                resume.format_run(desc_paragraph.add_run(desc), body_pt)
