from docx.shared import Pt

import resume


class Activity:
    def __init__(self, organization, location, role_title, description, delimiter):
        self.organization = organization
        self.location = location
        self.role_title = role_title#.split(delimiter)
        self.description = description.split(delimiter)

    def import_to_doc(self, doc, body_pt):
        """
        Imports this Activity to a given document
        :param doc: The document to import to
        :param body_pt: The size to make the body fields
        """
        organization_paragraph = doc.add_paragraph()
        organization_paragraph.paragraph_format.space_after = Pt(0)
        org_run = organization_paragraph.add_run(self.organization)
        resume.format_run(org_run, body_pt, bold=True)
        org_seperator = organization_paragraph.add_run(" | ")
        resume.format_run(org_seperator, body_pt)
        org_location = organization_paragraph.add_run(self.location)
        resume.format_run(org_location, body_pt, italic=True)

        role_title_paragraph = doc.add_paragraph()
        role_title_paragraph.paragraph_format.space_after = Pt(0)
        role_title_run = role_title_paragraph.add_run(self.role_title)
        resume.format_run(role_title_run, body_pt)

        if len(self.description) != 0:
            for desc in self.description:
                desc_paragraph = doc.add_paragraph()
                desc_paragraph.style = 'List Bullet'
                desc_paragraph.paragraph_format.space_after = Pt(0)
                resume.format_run(desc_paragraph.add_run(desc), body_pt)
