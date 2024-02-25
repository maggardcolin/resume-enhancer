from docx.shared import Pt

import resume


class Education:
    def __init__(self, degree, date, college, location, gpa, description, delimiter):
        self.degree = degree
        self.date = date
        self.college = college
        self.location = location
        self.gpa = gpa
        self.description = description.split(delimiter)

    def import_to_doc(self, doc, body_pt):
        """
        Imports this Education to a given document
        :param doc: The document to import to
        :param body_pt: The size to make the body fields
        """
        degree = self.degree + ", " + self.date
        degree_paragraph = doc.add_paragraph()
        degree_paragraph.paragraph_format.space_after = Pt(0)
        degree_run = degree_paragraph.add_run(degree)
        resume.format_run(degree_run, body_pt, bold=True)

        school_name_paragraph = doc.add_paragraph()
        school_name_paragraph.paragraph_format.space_after = Pt(0)
        school_name_run = school_name_paragraph.add_run(self.college + " | ")
        resume.format_run(school_name_run, body_pt)

        resume.format_run(school_name_paragraph.add_run(self.location), body_pt, italic=True)

        gpa_paragraph = doc.add_paragraph()
        gpa_paragraph.paragraph_format.space_after = Pt(0)
        gpa_run = gpa_paragraph.add_run("GPA: " + self.gpa)
        resume.format_run(gpa_run, body_pt)

        if len(self.description) != 0:
            for desc in self.description:
                desc_paragraph = doc.add_paragraph()
                desc_paragraph.style = 'List Bullet'
                desc_paragraph.paragraph_format.space_after = Pt(0)
                resume.format_run(desc_paragraph.add_run(desc), body_pt)
